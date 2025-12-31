# backend/routers/ai_writer.py

from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import FileResponse
from pydantic import BaseModel
import requests, json, re, os, uuid
from typing import List, Optional
from vector_db.client import get_collection

# PDF & DOCX export
from reportlab.pdfgen import canvas
from docx import Document

router = APIRouter(prefix="/ai-writer", tags=["ai-writer"])

OLLAMA_URL = "http://localhost:11434/api/generate"
MODEL_NAME = "gemma3:4b"

# ==================== Internal Utils ==================== #

def call_llm(prompt:str):
    """Unified LLM call with JSON auto-detection"""
    try:
        res = requests.post(OLLAMA_URL, json={"model":MODEL_NAME,"prompt":prompt,"stream":False}).json()
        raw = res.get("response","")
        block = re.search(r'\{.*\}|\[.*\]', raw, re.DOTALL)
        return json.loads(block.group(0)) if block else {"raw_output":raw}
    except Exception as e:
        raise HTTPException(500, f"LLM Failed → {e}")

def rag_context(limit=6000):
    """Retrieve stored research PDFs text from vector DB"""
    coll=get_collection().get(include=["documents","metadatas"])
    if not coll["documents"]: return "",[]
    return " ".join(coll["documents"])[:limit],[m.get("source","paper") for m in coll["metadatas"][:5]]

# ==================== Request Models ==================== #

class OutlineRequest(BaseModel):
    topic:str
    depth:str="medium"

class SectionRequest(BaseModel):
    topic:str
    section_title:str
    words:int=500
    use_docs:bool=True

class FullPaperRequest(BaseModel):
    topic:str
    words_per_section:int=600
    use_docs:bool=True

class TextRequest(BaseModel):
    text:str

class ExportRequest(BaseModel):
    title:str
    content:str

# ==================== 1. Outline ==================== #

@router.post("/outline")
def outline(req:OutlineRequest):
    prompt=f"""
Generate a research outline for topic: {req.topic}
Depth: {req.depth}
Return JSON: {{"outline":["Intro","Lit Review","Methodology","Results","Discussion","Conclusion"]}}
"""
    return call_llm(prompt)

# ==================== 2. Section Writer ==================== #

@router.post("/section")
def section(req:SectionRequest):
    ctx,sources = rag_context() if req.use_docs else ("",[])
    prompt=f"""
Write a research section:
Topic: {req.topic}
Section: {req.section_title}
Word target: {req.words}

Context (use if helpful):
{ctx}

Rules:
- 3-6 paragraphs
- add research gap insight
- include [CIT-1],[CIT-2] placeholders
Return JSON:
{{"title":"{req.section_title}","content":"...","citations":{sources}}}
"""
    return call_llm(prompt)

# ==================== 3. Full Paper ==================== #

@router.post("/full-paper")
def full_paper(req:FullPaperRequest):
    out=outline(OutlineRequest(topic=req.topic))
    sections = out.get("outline",[])
    
    final=""; all_cites=[]
    for sec in sections:
        part = section(SectionRequest(topic=req.topic,section_title=sec,words=req.words_per_section,use_docs=req.use_docs))
        final+=f"\n\n## {sec}\n{part['content']}"
        all_cites+=part.get("citations",[])

    abs = refine(TextRequest(text=final[:800]))
    keys = keywords(TextRequest(text=final))

    return {
        "title":req.topic,
        "sections":sections,
        "abstract":abs.get("refined",""),
        "keywords":keys.get("keywords",[]),
        "paper":final,
        "citations":list(set(all_cites))
    }

# ==================== 4. Refinement Tools ==================== #

@router.post("/refine")
def refine(req:TextRequest):
    prompt=f"Improve academically. Return JSON {{'refined':'text'}}\n{req.text}"
    return call_llm(prompt)

@router.post("/expand")
def expand(req:TextRequest):
    prompt=f"Expand academically. Return JSON {{'expanded':'text'}}\n{req.text}"
    return call_llm(prompt)

@router.post("/keywords")
def keywords(req:TextRequest):
    prompt=f"Extract 6-12 keywords. JSON {{'keywords':['..']}}\n{req.text}"
    return call_llm(prompt)

@router.post("/abstract")
def abstract(req:TextRequest):
    prompt=f"Create abstract from text. JSON {{'abstract':'...'}}\n{req.text}"
    return call_llm(prompt)

@router.post("/conclusion")
def conclusion(req:TextRequest):
    prompt=f"Write conclusion. JSON {{'conclusion':'...'}}\n{req.text}"
    return call_llm(prompt)

# ==================== 5. PDF Export ==================== #

@router.post("/export/pdf")
def export_pdf(req:ExportRequest):
    file=f"/tmp/{uuid.uuid4()}.pdf"
    c=canvas.Canvas(file)
    y=800
    for line in req.content.split("\n"):
        c.drawString(40,y,line[:120])
        y-=20
        if y<40:
            c.showPage(); y=800
    c.save()
    return FileResponse(file, filename=f"{req.title}.pdf")

# ==================== 6. DOCX Export ==================== #

@router.post("/export/docx")
def export_docx(req:ExportRequest):
    file=f"/tmp/{uuid.uuid4()}.docx"
    doc=Document()
    doc.add_heading(req.title, level=1)
    for p in req.content.split("\n"):
        doc.add_paragraph(p)
    doc.save(file)
    return FileResponse(file, filename=f"{req.title}.docx")
